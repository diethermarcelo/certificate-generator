import docx, json, logging

def create_doc():
    template_path = './template'
    template_filename = 'template.docx'
    template_fullpath = f"{template_path}/{template_filename}"
    doc = docx.Document(template_fullpath)
    return doc

def replace_string(data, p):
    terms = data.keys()
    for term in terms:
        # print('goes about here?')
        # print(p.text, 'ptext')
        if term in p.text:
            inline = p.runs
            # print('goes here')
            for i in range(len(inline)):
                replace_term = "{{" + term + "}}"
                text = inline[i].text.replace(replace_term, data[term])
                inline[i].text = text
def format_string(data):
    doc = create_doc()

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_string(data, paragraph)

    for p in doc.paragraphs:
        # print(terms, 'terms')
        # print(p.text, 'p.text')
        replace_string(data, p)

    return doc
def save_doc(doc, data, increment):
    current_file_name = data['file_name'] or "dest-{inc}".format(inc=increment)
    file_name = "./dist/{current_file_name}.docx".format(current_file_name=current_file_name)
    print("Creating {file_name}...".format(file_name=file_name))
    doc.save(file_name)

def main():
    try:
        json_file = open('./sample.json')
        full_data = json.load(json_file)
        i = 0
        for data in full_data:
            doc = format_string(data)
            save_doc(doc, data, i)
            i += 1
    except KeyError:
        print('please input a file_name attribute on your sample.json file')
    except:
        print('problem occurred. please try again later')
        logging.exception("message")
    return 1

if __name__ == '__main__':
    main()




